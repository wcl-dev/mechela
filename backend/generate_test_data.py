"""Generate anonymized, shortened test DOCX files for signal detection quality testing."""
from docx import Document
from pathlib import Path

OUT_DIR = Path(__file__).parent.parent / "demo"
OUT_DIR.mkdir(exist_ok=True)


def make_doc(filename, paragraphs):
    doc = Document()
    for style, text in paragraphs:
        if style == "heading":
            doc.add_heading(text, level=2)
        else:
            doc.add_paragraph(text)
    doc.save(OUT_DIR / filename)
    print(f"Created {filename} ({len(paragraphs)} paragraphs)")


# === Report 1: English Quarterly Report ===
# Mix of activities (non-signals) and L1/L2/L3 signal paragraphs
make_doc("demo_quarterly_report_en.docx", [
    ("heading", "Quarterly Progress Report"),
    ("heading", "Background"),
    ("normal", "The digital rights landscape in the Riverdale region has been shaped by rapid urbanization and uneven access to technology. Despite economic growth, many rural communities remain disconnected from meaningful participation in digital governance processes."),
    # L1 signal: enacted regulation
    ("normal", "A new national data protection regulation was enacted in March, requiring all telecommunications providers to publish transparency reports. This marks a significant shift in the regulatory environment."),
    ("normal", "Civic organizations have historically struggled to engage with technical standards bodies, partly due to a lack of accessible training materials and the highly specialized nature of these forums."),
    ("heading", "Objective 1: Strengthen civil society participation in digital governance"),
    ("heading", "Activities"),
    # Activity (non-signal)
    ("normal", "Greenfield Alliance attended the Regional Internet Governance Forum held in Mapleton. The team participated in three working sessions focused on content moderation and platform accountability."),
    # Activity (non-signal)
    ("normal", "We organized a two-day workshop on digital rights fundamentals, bringing together 28 participants from 12 local organizations. The workshop covered topics including data privacy, algorithmic transparency, and freedom of expression online."),
    # Activity (non-signal)
    ("normal", "The project coordinator submitted a written contribution to the open consultation on the proposed Digital Services Framework."),
    ("heading", "Outcomes"),
    # L2 signal: committed intent
    ("normal", "Following the workshop, four participating organizations formally committed to establishing a joint working group on digital governance. They signed a memorandum of understanding to coordinate their advocacy efforts over the next 18 months."),
    # L1 signal: institutional change
    ("normal", "Two legislators agreed to co-sponsor a parliamentary inquiry into algorithmic transparency in public services, after a series of briefing sessions organized by the coalition."),
    # L1 signal: policy adoption
    ("normal", "The Ministry of Communications adopted three of the five recommendations from our policy brief on rural connectivity, incorporating them into the revised National Broadband Plan."),
    # L3 signal: awareness
    ("normal", "Community awareness of digital rights issues has grown noticeably. A post-event survey showed that 73% of workshop attendees reported increased confidence in engaging with technical policy discussions."),
    ("heading", "Objective 2: Build a sustainable multi-stakeholder coalition"),
    # Activity (non-signal)
    ("normal", "We held recruitment meetings with potential coalition members from academia, the private sector, and grassroots organizations. A total of 15 new members joined the advisory council this quarter."),
    # L2 signal: intent
    ("normal", "The coalition drafted its first position paper on cross-border data flows, drawing on input from both technical experts and human rights scholars."),
    # L3 signal: interest
    ("normal", "Three regional telecom operators expressed interest in participating in a voluntary accountability pilot program proposed by the coalition."),
    ("heading", "Challenges"),
    # Non-signal
    ("normal", "Staff turnover at partner organizations slowed progress on the joint research initiative. Two key researchers left during the quarter, requiring us to revise the project timeline."),
    ("normal", "Securing meeting space for larger convenings remains difficult due to budget constraints. We are exploring partnerships with university facilities to address this."),
])


# === Report 2: English Final Narrative Report ===
# Heavier on signals (L1/L2) as a final report should be
make_doc("demo_final_report_en.docx", [
    ("heading", "Final Narrative Report"),
    ("heading", "Background"),
    ("normal", "The Sunridge Initiative was established to bridge the gap between human rights advocacy and technical standards development in the Pacific Coast region. The project operated from September 2023 to August 2024."),
    ("normal", "At the outset, civil society organizations in the region had limited capacity to engage with Internet governance processes. Technical discussions remained dominated by industry representatives and government agencies, with minimal input from rights-focused groups."),
    # Context signal: external environment
    ("normal", "Administrative fragmentation across multiple government agencies created inconsistent approaches to digital regulation. The telecommunications authority, the data protection commission, and the cybersecurity agency each operated under different mandates with little coordination."),
    ("heading", "Changes and Results"),
    # L1 signal: formal adoption
    ("normal", "Over the project period, the organization developed a comprehensive advocacy toolkit that has been formally adopted by six regional civil society groups as their standard framework for engaging with standards bodies."),
    # L1 signal: institutionalized
    ("normal", "We established the first regional Digital Rights Observatory, a permanent monitoring body that tracks implementation of technical standards affecting human rights. The Observatory now publishes quarterly assessments and has been formally recognized by the regional Internet governance forum as an official stakeholder group."),
    # L1 signal: policy incorporation
    ("normal", "Three policy recommendations from our research were incorporated into the revised National Cybersecurity Strategy, specifically regarding encryption standards, data localization requirements, and surveillance oversight mechanisms."),
    # L1 signal: institutionalized relationship
    ("normal", "The parliamentary committee on technology invited our executive director to serve as a permanent technical advisor on digital rights matters. This institutionalized relationship ensures ongoing civil society input into legislative processes."),
    # L2 signal: sustained shift
    ("normal", "The number of civil society representatives participating in regional standards meetings increased from 4 to 23 over the project period, representing a sustained shift in stakeholder composition."),
    ("heading", "Learning"),
    ("normal", "We initially assumed that many policies would be ready for immediate revision. However, the reality was that the region lacked a basic legal framework for Internet governance, requiring us to focus first on foundational capacity building before pursuing specific policy changes."),
    ("normal", "Building trust with legislators required more time than anticipated. We found that sustained engagement over multiple meetings was far more effective than one-off presentations. The most productive relationships developed through informal consultations rather than formal hearings."),
    ("normal", "The balance of technical and non-technical stakeholders in our convenings was consistently praised by participants. Maintaining this balance required deliberate design choices in agenda setting and facilitation methods."),
    ("normal", "Monitoring trends in international standards organizations proved highly valuable for anticipating domestic policy discussions. Several emerging issues at the global level arrived in regional debates 6-12 months later, giving our coalition time to prepare informed positions."),
    ("heading", "Sustainability"),
    # L2 signal: secured funding
    ("normal", "The coalition has secured independent funding to continue operations for the next two years. The advisory council continues to meet monthly, and three working groups remain active on encryption policy, platform governance, and accessibility standards."),
    # L1 signal: institutionalized into curricula
    ("normal", "Two universities have integrated materials from our training program into their graduate curricula on technology policy, ensuring knowledge transfer beyond the project period."),
])


# === Report 3: Second Quarterly Report ===
# More activity-heavy, fewer signals (typical early-stage report)
make_doc("demo_quarterly_report_q2.docx", [
    ("heading", "Quarterly Progress Report - Q2"),
    ("heading", "Executive Summary"),
    ("normal", "Harbor Foundation continued implementing the Digital Accountability Initiative during the second quarter. Key achievements include completing the localized research methodology, launching the stakeholder mapping exercise, and initiating company engagement."),
    ("heading", "Objective 1: Adapt international accountability standards for regional application"),
    ("heading", "Activities"),
    # Activity (non-signal)
    ("normal", "Harbor Foundation produced a detailed assessment of the Digital Governance Assessment Framework, identifying areas requiring adaptation for the regional regulatory context. The assessment was shared with all project partners."),
    # Activity (non-signal)
    ("normal", "We developed a training program on digital rights research methods. Two newly recruited researchers completed the program and began independent work on company profiles."),
    # Activity (non-signal)
    ("normal", "Harbor Foundation facilitated a coordination meeting with partner organizations to align on research timelines and share preliminary findings from desk research."),
    # Activity (non-signal)
    ("normal", "To maintain public awareness, we published three articles on corporate digital accountability in our newsletter, reaching approximately 4,500 subscribers."),
    ("heading", "Objective 2: Build regional capacity for corporate accountability research"),
    # Activity (non-signal)
    ("normal", "The research team completed a regulatory landscape analysis covering data protection, content moderation, and platform liability frameworks across five jurisdictions in the region."),
    # Activity (non-signal)
    ("normal", "We translated and adapted the full set of research indicators into the local language. This required extensive consultation with legal experts to ensure terminology accuracy."),
    # L2 signal: agreed to pilot
    ("normal", "Three partner organizations agreed to pilot the adapted methodology in their respective countries, forming a regional research network that will produce comparable findings."),
    # L3 signal: preliminary findings (awareness level)
    ("normal", "The preliminary findings suggest that most major telecommunications companies in the region have adopted basic transparency reporting practices, but significant gaps remain in areas such as algorithmic accountability and government request disclosures."),
    ("heading", "Challenges"),
    ("normal", "Understanding the precise definitions of some research indicators proved difficult due to the specialized legal language used in the original methodology. We addressed this through weekly discussion sessions with the methodology team."),
    ("normal", "Developing company selection criteria required more deliberation than expected. Current guidance provides limited direction for selecting companies in markets with different ownership structures and regulatory frameworks."),
    ("heading", "Planned Activities for Next Quarter"),
    ("normal", "Finalize the research plan including company selection, indicator prioritization, and data collection workflow."),
    ("normal", "Begin primary data collection through company questionnaires and stakeholder interviews."),
    ("normal", "Publish the adapted research methodology guide on our website to support transparency and potential replication by other organizations."),
])


print("Done. Files saved to demo/")
