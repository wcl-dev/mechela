"""
DOCX parser: extracts paragraphs and table cells,
detects section headings, and builds context windows.
"""
from pathlib import Path
from dataclasses import dataclass
import docx

SECTION_KEYWORDS = [
    "background", "context", "activities", "progress",
    "outcomes", "results", "challenges", "lessons",
    "objectives", "executive summary", "introduction",
]


@dataclass
class ParsedAnchor:
    paragraph_index: int
    section: str | None
    text: str
    context_text: str  # target + one before + one after


def _detect_section(text: str) -> str | None:
    lower = text.lower().strip()
    for kw in SECTION_KEYWORDS:
        if lower.startswith(kw) or kw in lower[:60]:
            return kw
    return None


def _is_heading(paragraph) -> bool:
    style_name = (paragraph.style.name or "").lower()
    return "heading" in style_name or paragraph.runs and all(
        r.bold for r in paragraph.runs if r.text.strip()
    )


def _clean(text: str) -> str:
    """Normalize surrogate-escaped bytes (e.g. Word em-dash) to proper Unicode."""
    return text.encode("utf-8", "surrogatepass").decode("utf-8")


def parse_docx(file_path: Path) -> list[ParsedAnchor]:
    doc = docx.Document(str(file_path))
    raw_texts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = _clean(para.text.strip())
        if text:
            raw_texts.append(text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = _clean(cell.text.strip())
                if text and text not in raw_texts:
                    raw_texts.append(text)

    if not raw_texts:
        return []

    # Assign sections and build context windows
    anchors: list[ParsedAnchor] = []
    current_section: str | None = None

    for i, text in enumerate(raw_texts):
        detected = _detect_section(text)
        if detected:
            current_section = detected

        # Context window: previous + current + next paragraph
        parts = []
        if i > 0:
            parts.append(raw_texts[i - 1])
        parts.append(text)
        if i < len(raw_texts) - 1:
            parts.append(raw_texts[i + 1])
        context_text = " | ".join(parts)

        anchors.append(ParsedAnchor(
            paragraph_index=i,
            section=current_section,
            text=text,
            context_text=context_text,
        ))

    return anchors
