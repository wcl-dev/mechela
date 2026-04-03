"""
Generate Mechela demo DOCX — Q3 2024 Quarterly Progress Report
Run: py generate_demo.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Styles ──────────────────────────────────────────────
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def spacer():
    doc.add_paragraph()

# ── Cover ────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Quarterly Progress Report")
run.bold = True
run.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Digital Inclusion & Rights Advocacy Programme\n")
sub.add_run("Q3 2024 | July – September 2024\n")
sub.add_run("Community Digital Rights Initiative (CDRI)")

spacer()
doc.add_paragraph("──────────────────────────────────────────")
spacer()

# ── 1. Background and Context ────────────────────────────
heading("1. Background and Context")

para(
    "The third quarter of 2024 unfolded against a rapidly shifting digital governance landscape "
    "across the region. In July 2024, the national government enacted the Personal Data Protection "
    "Act, establishing for the first time a statutory right to data privacy and a dedicated "
    "enforcement authority. This development significantly elevated the policy environment in which "
    "CDRI operates and created new entry points for civil society engagement."
)

para(
    "At the regional level, the Southeast Asia Internet Governance Forum adopted a non-binding "
    "resolution recognising internet access as a foundational enabler of human rights. Although "
    "non-binding, the resolution was endorsed by eleven member governments and signals a normative "
    "shift that advocacy organisations can leverage in national dialogues."
)

para(
    "In parallel, the regional telecommunications standards body initiated a public consultation "
    "on minimum service obligations for rural broadband, reflecting growing political pressure to "
    "address the urban-rural digital divide. CDRI submitted a formal response to this consultation "
    "in August."
)

spacer()

# ── 2. Programme Objectives ──────────────────────────────
heading("2. Programme Objectives")

para("This programme operates under three strategic objectives:")

doc.add_paragraph("Objective 1 — Policy Influence: Embed digital rights principles into national and regional governance frameworks.", style="List Number")
doc.add_paragraph("Objective 2 — Institutional Capacity: Strengthen civil society's ability to monitor, document, and respond to digital rights developments.", style="List Number")
doc.add_paragraph("Objective 3 — Community Awareness: Build understanding of digital rights among marginalised communities, with a focus on women and youth.", style="List Number")

spacer()

# ── 3. Key Achievements and Outcomes ────────────────────
heading("3. Key Achievements and Outcomes")

heading("3.1 Policy and Institutional Change", level=2)

para(
    "The quarter produced significant progress on Objective 1. Following eighteen months of "
    "sustained advocacy, the Ministry of Communications formally incorporated digital rights "
    "language — including explicit references to freedom of expression, privacy, and non-discrimination "
    "online — into the National Digital Strategy 2024–2030, which was officially gazetted in "
    "September 2024. This marks the first time rights-based framing has been institutionalised "
    "in a national digital policy instrument."
)

para(
    "At the sub-national level, four provincial governments signed a memorandum of understanding "
    "with CDRI to co-develop a Digital Rights Municipal Charter. The charter commits signatory "
    "municipalities to conduct annual digital rights impact assessments of proposed local ordinances "
    "and to establish a standing civil society advisory panel on digital governance."
)

para(
    "The national telecommunications regulator agreed to establish a permanent Civil Society "
    "Liaison Desk within its policy division, following CDRI's submission to the broadband "
    "consultation and a series of bilateral meetings held in August and September. The Liaison Desk "
    "is expected to be operational by December 2024."
)

heading("3.2 Capacity and Coalition Building", level=2)

para(
    "Under Objective 2, CDRI convened the first formal meeting of the Digital Rights Monitoring "
    "Network (DRMN) in August 2024, bringing together fourteen civil society organisations from "
    "eight provinces. The network adopted a shared monitoring methodology and agreed to produce "
    "a joint quarterly shadow report beginning in Q4 2024."
)

para(
    "The Ministry of Education expressed strong interest in integrating CDRI's Digital Citizenship "
    "curriculum into the secondary school syllabus review scheduled for early 2025. Preliminary "
    "discussions have been held with the Curriculum Development Bureau, and a formal proposal "
    "is under preparation."
)

para(
    "Three partner organisations that participated in CDRI's capacity-building programme have "
    "independently begun using structured evidence documentation in their own M&E reporting, "
    "indicating early uptake of practices introduced through the programme."
)

heading("3.3 Community Awareness", level=2)

para(
    "Under Objective 3, community awareness sessions reached 1,840 participants across six "
    "provinces, exceeding the Q3 target of 1,500. Post-session surveys indicate that 78% of "
    "participants reported increased confidence in identifying online rights violations, compared "
    "to 41% in the baseline."
)

para(
    "A community paralegal network in the Northern Region, trained under a previous CDRI cohort, "
    "successfully handled its first digital rights complaint — a case of non-consensual image "
    "sharing — through the formal legal aid system, resulting in a mediated resolution and a "
    "written undertaking from the respondent."
)

para(
    "Local women's groups in two districts have begun independently facilitating peer digital "
    "safety sessions, using materials adapted from CDRI's training kit. This peer-led model was "
    "not part of the original programme design but emerged organically from participant initiative."
)

spacer()

# ── 4. Activities Summary ────────────────────────────────
heading("4. Activities Summary")

table = doc.add_table(rows=1, cols=4)
table.style = "Table Grid"

hdr = table.rows[0].cells
hdr[0].text = "Activity"
hdr[1].text = "Target"
hdr[2].text = "Achieved"
hdr[3].text = "Notes"

rows = [
    ("Policy advocacy meetings", "12", "15", "Includes 3 unplanned bilateral meetings with regulator"),
    ("Community awareness sessions", "20", "23", "Exceeded due to partner co-facilitation"),
    ("Civil society capacity workshops", "4", "4", "On track"),
    ("Research and publications", "2", "2", "Joint shadow report methodology paper; broadband consultation submission"),
    ("Coalition coordination meetings", "3", "4", "Additional session due to DRMN launch"),
]

for activity, target, achieved, notes in rows:
    row = table.add_row().cells
    row[0].text = activity
    row[1].text = target
    row[2].text = achieved
    row[3].text = notes

spacer()

# ── 5. Challenges and Lessons Learned ───────────────────
heading("5. Challenges and Lessons Learned")

para(
    "Staff turnover in two partner organisations affected the pace of coalition coordination "
    "during July and August. CDRI responded by increasing direct technical support to affected "
    "partners, which partially offset the disruption but added unplanned workload to programme staff."
)

para(
    "Community sessions in rural areas continue to face logistical barriers, including unreliable "
    "transport and limited availability of suitable venues. Co-facilitation with local women's groups "
    "has proven an effective mitigation strategy and will be formalised in the Q4 workplan."
)

para(
    "The pace of policy engagement was faster than anticipated following the enactment of the "
    "Data Protection Act, requiring the team to respond to multiple simultaneous consultation "
    "processes. This has surfaced a need for clearer internal prioritisation criteria for advocacy "
    "opportunities."
)

spacer()

# ── 6. Next Steps ────────────────────────────────────────
heading("6. Next Steps (Q4 2024)")

doc.add_paragraph("Finalise and submit the Digital Rights Municipal Charter to the four signatory municipalities.", style="List Bullet")
doc.add_paragraph("Support the telecommunications regulator in establishing the Civil Society Liaison Desk.", style="List Bullet")
doc.add_paragraph("Produce the first joint DRMN quarterly shadow report.", style="List Bullet")
doc.add_paragraph("Submit formal curriculum integration proposal to the Ministry of Education.", style="List Bullet")
doc.add_paragraph("Conduct end-of-year evaluation of community awareness sessions.", style="List Bullet")

spacer()
doc.add_paragraph("──────────────────────────────────────────")
para("End of Report", bold=True)
para("Community Digital Rights Initiative (CDRI) | Q3 2024 Quarterly Progress Report")

# ── Save ─────────────────────────────────────────────────
output = "CDRI_Q3_2024_Progress_Report.docx"
doc.save(output)
print(f"Saved: {output}")
